"""
试卷导出工具
作者：智学伴开发团队
目的：支持试卷导出为PDF/Word格式
"""
import os
import platform
import re
import tempfile
import base64
from typing import Dict, List, Any, Tuple, Optional
from datetime import datetime
from xml.sax.saxutils import escape as xml_escape
from core.logger import logger

# 尝试导入reportlab（PDF导出）
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image as RLImage
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger.warning("reportlab未安装，PDF导出功能不可用。请运行: pip install reportlab")

# 尝试导入python-docx（Word导出）
DOCX_AVAILABLE = False
DOCX_ERROR_MSG = None
try:
    # 先检查lxml依赖
    try:
        from lxml import etree
    except ImportError:
        DOCX_AVAILABLE = False
        DOCX_ERROR_MSG = "python-docx依赖lxml库，但lxml未正确安装。请运行: pip install --upgrade lxml"
        logger.warning(DOCX_ERROR_MSG)
    else:
        # lxml可用，尝试导入python-docx
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    DOCX_ERROR_MSG = f"python-docx未正确安装: {e}。请运行: pip install python-docx"
    logger.warning(DOCX_ERROR_MSG)
except Exception as e:
    # 处理其他导入错误（如安装了错误的docx包或lxml版本不兼容）
    DOCX_AVAILABLE = False
    error_msg = str(e)
    if "lxml" in error_msg or "etree" in error_msg:
        DOCX_ERROR_MSG = f"python-docx依赖lxml库，但lxml导入失败: {e}。请运行: pip install --upgrade lxml"
    else:
        DOCX_ERROR_MSG = f"python-docx导入失败: {e}。如果安装了错误的docx包，请运行: pip uninstall docx && pip install python-docx"
    logger.warning(DOCX_ERROR_MSG)

# 全局字体注册标志
_registered_font_name = None


class PaperExporter:
    @staticmethod
    def _normalize_export_questions(questions: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        normalized_questions: List[Dict[str, Any]] = []
        for question in questions or []:
            stem = question.get("question") or question.get("stem") or ""
            normalized_questions.append(
                {
                    **question,
                    "question": stem,
                    "stem": question.get("stem") or stem,
                    "points": max(1, int(question.get("points") or 1)),
                    "question_images": question.get("question_images") or [],
                    "answer_images": question.get("answer_images") or [],
                    "solution_images": question.get("solution_images") or [],
                }
            )
        return normalized_questions

    @staticmethod
    def _normalize_answer_key(answer_key: Any) -> Dict[str, Dict[str, Any]]:
        if isinstance(answer_key, dict):
            return {
                str(key): value
                for key, value in answer_key.items()
                if isinstance(value, dict)
            }
        if isinstance(answer_key, list):
            normalized: Dict[str, Dict[str, Any]] = {}
            for index, item in enumerate(answer_key, start=1):
                if isinstance(item, dict):
                    normalized[str(index)] = item
                elif isinstance(item, str):
                    normalized[str(index)] = {"answer": item}
            return normalized
        return {}

    @staticmethod
    def _resolve_asset_path(asset: Any) -> Optional[str]:
        if not asset:
            return None
        if isinstance(asset, str):
            candidate = asset
        elif isinstance(asset, dict):
            candidate = (
                asset.get("file_path")
                or asset.get("local_path")
                or asset.get("path")
                or asset.get("preview_url")
                or asset.get("image_url")
                or asset.get("file_url")
            )
        else:
            candidate = str(asset)

        if not candidate:
            return None

        if candidate.startswith(("http://", "https://", "data:")):
            return None

        normalized = candidate.lstrip("/")
        upload_relative = normalized[len("uploads/"):] if normalized.startswith("uploads/") else normalized
        upload_candidate = os.path.join("uploads", upload_relative)
        if os.path.isabs(candidate) and os.path.exists(candidate):
            return candidate
        if os.path.exists(candidate):
            return candidate
        if os.path.exists(normalized):
            return normalized
        if os.path.exists(upload_candidate):
            return upload_candidate
        return None

    @staticmethod
    def _append_images_to_pdf_story(story: List[Any], assets: List[Any], max_width: float = 14 * cm, max_height: float = 8 * cm) -> None:
        for asset in assets or []:
            image_path = PaperExporter._resolve_asset_path(asset)
            if not image_path:
                continue
            try:
                from PIL import Image as PILImage
                with PILImage.open(image_path) as pil_img:
                    img_w, img_h = pil_img.size
                scale = min(max_width / max(img_w, 1), max_height / max(img_h, 1), 1)
                story.append(RLImage(image_path, width=img_w * scale, height=img_h * scale))
                story.append(Spacer(1, 0.3 * cm))
            except Exception as exc:
                logger.warning("Failed to render paper image in PDF export: %s", exc)

    @staticmethod
    def _append_images_to_word_doc(doc, assets: List[Any], width_inches: float = 4.8) -> None:
        for asset in assets or []:
            image_path = PaperExporter._resolve_asset_path(asset)
            if not image_path:
                continue
            try:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
            except Exception as exc:
                logger.warning("Failed to render paper image in Word export: %s", exc)
    """试卷导出类"""
    
    @staticmethod
    def _render_latex_to_image(latex_formula: str, output_path: Optional[str] = None) -> Optional[str]:
        """
        将LaTeX公式渲染为图片
        
        Args:
            latex_formula: LaTeX公式字符串
            output_path: 输出图片路径（可选，如果不提供则创建临时文件）
            
        Returns:
            str: 图片文件路径，如果渲染失败返回None
        """
        try:
            import matplotlib
            matplotlib.use('Agg')  # 使用非交互式后端
            import matplotlib.pyplot as plt
            from matplotlib import mathtext
            
            # 清理LaTeX公式
            latex_formula = latex_formula.strip()
            # 移除可能存在的公式定界符
            if latex_formula.startswith('$$') and latex_formula.endswith('$$'):
                latex_formula = latex_formula[2:-2]
            elif latex_formula.startswith('$') and latex_formula.endswith('$'):
                latex_formula = latex_formula[1:-1]
            elif latex_formula.startswith('\\[') and latex_formula.endswith('\\]'):
                latex_formula = latex_formula[2:-2]
            elif latex_formula.startswith('\\(') and latex_formula.endswith('\\)'):
                latex_formula = latex_formula[2:-2]
            
            # 创建临时文件（如果没有提供路径）
            if not output_path:
                temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                output_path = temp_file.name
                temp_file.close()
            
            # 创建图形
            fig, ax = plt.subplots(figsize=(8, 2))
            ax.axis('off')
            
            # 渲染LaTeX公式
            try:
                # 尝试使用matplotlib的数学文本渲染器（不需要系统LaTeX）
                ax.text(0.5, 0.5, f'${latex_formula}$', 
                       fontsize=16, ha='center', va='center',
                       transform=ax.transAxes,
                       usetex=False)  # 不使用系统LaTeX，使用matplotlib内置渲染器
            except Exception as e:
                logger.warning(f"LaTeX渲染失败，尝试使用文本模式: {e}")
                # 如果LaTeX渲染失败，使用文本模式
                ax.text(0.5, 0.5, latex_formula,
                       fontsize=16, ha='center', va='center',
                       transform=ax.transAxes)
            
            # 保存图片
            plt.savefig(output_path, bbox_inches='tight', dpi=150, 
                       facecolor='white', edgecolor='none', pad_inches=0.1)
            plt.close(fig)
            
            return output_path
        except ImportError:
            logger.warning("matplotlib未安装，无法渲染LaTeX公式为图片")
            return None
        except Exception as e:
            logger.warning(f"LaTeX公式渲染失败: {e}")
            return None
    
    @staticmethod
    def _convert_latex_to_word_math(text: str) -> List[Tuple[str, bool, Optional[str]]]:
        """
        将包含LaTeX公式的文本转换为Word可用的格式
        返回一个列表，每个元素是(文本片段, 是否为LaTeX公式)
        
        Args:
            text: 包含LaTeX公式的文本
            
        Returns:
            List[Tuple[str, bool, Optional[str]]]: [(文本片段, 是否为LaTeX, 图片路径)]
        """
        if not text:
            return [("", False, None)]
        
        try:
            # 匹配LaTeX公式的模式：$...$、\(...\)、\[...\]
            patterns = [
                (r'\$\$([^$]+)\$\$', True),  # 块级公式 $$...$$
                (r'\$([^$]+)\$', True),      # 行内公式 $...$
                (r'\\\[([^\]]+)\\\]', True), # 块级公式 \[...\]
                (r'\\\(([^\)]+)\\\)', True), # 行内公式 \(...\)
            ]
            
            result = []
            last_pos = 0
            
            # 找到所有匹配的LaTeX公式
            matches = []
            for pattern, is_math in patterns:
                for match in re.finditer(pattern, text):
                    matches.append((match.start(), match.end(), match.group(1), is_math))
            
            # 按位置排序
            matches.sort(key=lambda x: x[0])
            
            # 处理重叠的匹配（保留最长的）
            filtered_matches = []
            for i, (start, end, content, is_math) in enumerate(matches):
                # 检查是否与已有匹配重叠
                overlap = False
                for prev_start, prev_end, _, _ in filtered_matches:
                    if not (end <= prev_start or start >= prev_end):
                        overlap = True
                        # 如果当前匹配更长，替换之前的
                        if end - start > prev_end - prev_start:
                            filtered_matches = [(s, e, c, m) for s, e, c, m in filtered_matches 
                                              if not (s == prev_start and e == prev_end)]
                            overlap = False
                        break
                if not overlap:
                    filtered_matches.append((start, end, content, is_math))
            
            # 按位置重新排序
            filtered_matches.sort(key=lambda x: x[0])
            
            # 构建结果列表
            for start, end, content, is_math in filtered_matches:
                # 添加公式前的普通文本
                if start > last_pos:
                    result.append((text[last_pos:start], False, None))
                # 尝试将LaTeX公式渲染为图片
                image_path = PaperExporter._render_latex_to_image(content.strip())
                if image_path:
                    # 如果渲染成功，使用图片
                    result.append((content.strip(), True, image_path))
                else:
                    # 如果渲染失败，转换为简单的数学表示
                    math_text = content.strip()
                    # 简单的LaTeX到文本的转换
                    math_text = math_text.replace('\\frac{', '(').replace('}{', '/').replace('}', ')')
                    math_text = math_text.replace('\\sqrt{', '√(').replace('}', ')')
                    math_text = math_text.replace('^', '^').replace('_', '_')
                    math_text = math_text.replace('\\lim', 'lim').replace('\\sin', 'sin')
                    math_text = math_text.replace('\\cos', 'cos').replace('\\tan', 'tan')
                    math_text = math_text.replace('\\pi', 'π').replace('\\infty', '∞')
                    math_text = math_text.replace('\\alpha', 'α').replace('\\beta', 'β')
                    math_text = math_text.replace('\\gamma', 'γ').replace('\\theta', 'θ')
                    result.append((math_text, True, None))
                last_pos = end
            
            # 添加剩余的普通文本
            if last_pos < len(text):
                result.append((text[last_pos:], False, None))
            
            # 如果没有匹配到任何公式，返回原文本
            if not result:
                result = [(text, False, None)]
            
            return result
        except Exception as e:
            # 如果处理失败，返回原文本
            logger.warning(f"LaTeX转换失败，使用原文本: {e}")
            return [(text, False, None)]

    @staticmethod
    def _detect_latex_in_text(text: str) -> List[Tuple[str, bool]]:
        """Detect LaTeX formulas in text and return segments.

        Returns list of (segment_text, is_latex) tuples.
        """
        if not text:
            return [("", False)]

        segments: List[Tuple[str, bool]] = []
        pattern = re.compile(
            r'(\$\$[\s\S]+?\$\$|\$[^$\n]+\$|\\\[[\s\S]+?\\\]|\\\([\s\S]+?\\\)|\\begin\{[A-Za-z*]+\}[\s\S]+?\\end\{[A-Za-z*]+\})'
        )

        parts = pattern.split(text)
        for part in parts:
            if not part:
                continue
            stripped = part.strip()
            is_latex = bool(
                (stripped.startswith("$") and stripped.endswith("$"))
                or (stripped.startswith("\\[") and stripped.endswith("\\]"))
                or (stripped.startswith("\\(") and stripped.endswith("\\)"))
                or stripped.startswith("\\begin{")
            )
            segments.append((part, is_latex))

        if not segments:
            segments = [(text, False)]
        return segments

    @staticmethod
    def _add_text_with_latex_to_story(
        story,
        text,
        para_style,
        chinese_font,
        image_max_width=14 * cm,
        prefix_markup: str = "",
        suffix_markup: str = "",
    ):
        """Add text (possibly containing LaTeX) to a ReportLab story list.

        Text segments become Paragraphs; LaTeX formulas are rendered as images
        via _render_latex_to_image and inserted as Image flowables.
        """
        if text is None:
            text = ""
        if not str(text).strip() and (prefix_markup or suffix_markup):
            story.append(Paragraph(f"{prefix_markup}{suffix_markup}", para_style))
            return
        if not str(text).strip():
            return

        segments = [
            (segment_text, is_latex)
            for segment_text, is_latex in PaperExporter._detect_latex_in_text(str(text))
            if segment_text.strip()
        ]
        if not segments:
            if prefix_markup or suffix_markup:
                story.append(Paragraph(f"{prefix_markup}{suffix_markup}", para_style))
            return

        prefix_pending = prefix_markup
        suffix_pending = suffix_markup
        for index, (segment_text, is_latex) in enumerate(segments):
            if not segment_text.strip():
                continue
            has_later_text = any(not later_is_latex and later_text.strip() for later_text, later_is_latex in segments[index + 1:])

            if is_latex:
                if prefix_pending:
                    story.append(Paragraph(prefix_pending, para_style))
                    prefix_pending = ""
                image_path = PaperExporter._render_latex_to_image(segment_text)
                if image_path:
                    try:
                        from reportlab.platypus import Image
                        from PIL import Image as PILImage
                        with PILImage.open(image_path) as pil_img:
                            img_w, img_h = pil_img.size
                        aspect = img_h / max(img_w, 1)
                        target_w = min(image_max_width, img_w)
                        target_h = target_w * aspect
                        story.append(Image(image_path, width=target_w, height=target_h))
                        story.append(Spacer(1, 0.15 * cm))
                    except Exception:
                        fallback = PaperExporter._latex_fallback_text(segment_text)
                        if fallback:
                            story.append(Paragraph(xml_escape(fallback), para_style))
                else:
                    fallback = PaperExporter._latex_fallback_text(segment_text)
                    if fallback:
                        story.append(Paragraph(xml_escape(fallback), para_style))
                if suffix_pending and not has_later_text:
                    story.append(Paragraph(suffix_pending, para_style))
                    suffix_pending = ""
            else:
                escaped_text = xml_escape(segment_text).replace("\n", "<br/>")
                paragraph_text = f"{prefix_pending}{escaped_text}"
                prefix_pending = ""
                if suffix_pending and not has_later_text:
                    paragraph_text = f"{paragraph_text}{suffix_pending}"
                    suffix_pending = ""
                story.append(Paragraph(paragraph_text, para_style))

        if suffix_pending:
            story.append(Paragraph(suffix_pending, para_style))

    @staticmethod
    def _latex_fallback_text(latex_str: str) -> str:
        """Convert simple LaTeX to readable plain text when rendering fails."""
        text = latex_str.strip()
        for dollar in ("$$", "$"):
            if text.startswith(dollar) and text.endswith(dollar):
                text = text[len(dollar):-len(dollar)].strip()
                break
        for bracket in (("\\[", "\\]"), ("\\(", "\\)")):
            if text.startswith(bracket[0]) and text.endswith(bracket[1]):
                text = text[len(bracket[0]):-len(bracket[1])].strip()
                break
        text = re.sub(r"\\begin\{([A-Za-z*]+)\}", r"\1: ", text)
        text = re.sub(r"\\end\{[A-Za-z*]+\}", "", text)

        replacements = {
            "\\frac": "", "\\sqrt": "√", "\\sum": "Σ", "\\int": "∫",
            "\\pi": "π", "\\infty": "∞", "\\alpha": "α", "\\beta": "β",
            "\\gamma": "γ", "\\theta": "θ", "\\lambda": "λ", "\\mu": "μ",
            "\\sigma": "σ", "\\omega": "ω", "\\Delta": "Δ", "\\times": "×",
            "\\div": "÷", "\\pm": "±", "\\cdot": "·", "\\leq": "≤",
            "\\geq": "≥", "\\neq": "≠", "\\approx": "≈", "\\rightarrow": "→",
            "\\Rightarrow": "⇒", "\\sin": "sin", "\\cos": "cos", "\\tan": "tan",
            "\\log": "log", "\\ln": "ln", "\\lim": "lim",
        }
        for latex, plain in replacements.items():
            text = text.replace(latex, plain)
        text = re.sub(r"[{}_^]", "", text)
        return text.strip()
    
    @staticmethod
    def _register_chinese_font():
        """注册中文字体"""
        global _registered_font_name
        
        if _registered_font_name and _registered_font_name in pdfmetrics.getRegisteredFontNames():
            return _registered_font_name
        
        system = platform.system()
        font_paths = []
        
        if system == "Windows":
            font_paths = [
                ("SimHei", "C:/Windows/Fonts/simhei.ttf"),  # 黑体
                ("SimSun", "C:/Windows/Fonts/simsun.ttc"),  # 宋体
                ("SimKai", "C:/Windows/Fonts/simkai.ttf"),  # 楷体
                ("Microsoft YaHei", "C:/Windows/Fonts/msyh.ttc"),  # 微软雅黑
            ]
        elif system == "Darwin":  # macOS
            font_paths = [
                ("STHeiti", "/System/Library/Fonts/STHeiti Light.ttc"),
                ("PingFang SC", "/System/Library/Fonts/PingFang.ttc"),
            ]
        else:  # Linux
            font_paths = [
                ("WenQuanYi Micro Hei", "/usr/share/fonts/wqy-microhei/wqy-microhei.ttc"),
                ("WenQuanYi Zen Hei", "/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc"),
            ]
        
        registered_font = None
        
        # 尝试注册系统字体
        for font_name, font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    if font_path.lower().endswith('.ttc'):
                        # TTC文件是字体集合，尝试直接加载
                        font = TTFont(font_name, font_path)
                        pdfmetrics.registerFont(font)
                        registered_font = font_name
                        logger.info(f"✅ 成功注册中文字体: {font_name}")
                        break
                    else:
                        # TTF文件直接注册
                        font = TTFont(font_name, font_path)
                        pdfmetrics.registerFont(font)
                        registered_font = font_name
                        logger.info(f"✅ 成功注册中文字体: {font_name}")
                        break
                except Exception as e:
                    logger.debug(f"注册字体失败 {font_name}: {e}")
                    continue
        
        # 如果没有找到系统字体，尝试使用reportlab内置的CID字体
        if not registered_font:
            try:
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
                registered_font = "STSong-Light"
                logger.info("✅ 使用ReportLab内置中文字体: STSong-Light")
            except Exception as e:
                logger.warning(f"无法注册CID字体: {e}")
                registered_font = "Helvetica"  # 回退到默认字体
        
        if not registered_font:
            registered_font = "Helvetica"
        
        _registered_font_name = registered_font
        return registered_font
    
    @staticmethod
    def export_to_pdf(paper_data: Dict[str, Any], output_path: str) -> str:
        """
        导出试卷为PDF
        
        Args:
            paper_data: 试卷数据
                - title: 试卷标题
                - subject: 科目
                - grade_level: 学段
                - total_questions: 总题数
                - time_limit: 考试时长
                - questions: 题目列表
                - answer_key: 标准答案（可选）
            output_path: 输出文件路径
        
        Returns:
            str: 输出文件路径
        """
        if not REPORTLAB_AVAILABLE:
            raise ValueError("PDF导出功能需要安装reportlab库：pip install reportlab")
        
        try:
            # 注册中文字体
            chinese_font = PaperExporter._register_chinese_font()
            
            # 创建PDF文档
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=2*cm,
                leftMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )
            
            # 构建内容
            story = []
            styles = getSampleStyleSheet()
            
            # 标题样式（使用中文字体）
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=chinese_font,
                fontSize=18,
                textColor=colors.HexColor('#1e40af'),
                spaceAfter=30,
                alignment=1  # 居中
            )
            
            # 正文样式（使用中文字体）
            normal_style = ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=chinese_font,
                fontSize=12,
                leading=18
            )
            
            # 添加标题
            title = paper_data.get("title", "试卷")
            story.append(Paragraph(xml_escape(str(title)), title_style))
            story.append(Spacer(1, 0.5*cm))
            
            # 添加试卷信息
            info_data = []
            if paper_data.get("subject"):
                info_data.append(["科目", paper_data["subject"]])
            if paper_data.get("grade_level"):
                info_data.append(["学段", paper_data["grade_level"]])
            if paper_data.get("total_questions"):
                info_data.append(["题数", str(paper_data["total_questions"])])
            if paper_data.get("time_limit"):
                info_data.append(["时长", f"{paper_data['time_limit']}分钟"])
            
            if info_data:
                info_table = Table(info_data, colWidths=[3*cm, 6*cm])
                info_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), chinese_font),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                    ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                story.append(info_table)
                story.append(Spacer(1, 1*cm))
            
            # 添加题目
            questions = PaperExporter._normalize_export_questions(paper_data.get("questions", []))
            for i, q in enumerate(questions, 1):
                q_type = q.get('type', '')
                type_label = ''
                if q_type == 'choice':
                    type_label = '单选题'
                elif q_type == 'multiple_choice':
                    type_label = '多选题'
                elif q_type == 'fill':
                    type_label = '填空题'
                elif q_type == 'judge':
                    type_label = '判断题'
                elif q_type == 'essay':
                    type_label = '简答题'
                elif q_type == 'calculation':
                    type_label = '计算题'

                question_suffix = ""
                if type_label:
                    question_suffix += f" <i>（{xml_escape(type_label)}）</i>"
                question_suffix += f" <i>（{q.get('points', 5)}分）</i>"
                PaperExporter._add_text_with_latex_to_story(
                    story,
                    q.get("question", ""),
                    normal_style,
                    chinese_font,
                    prefix_markup=f"<b>{i}.</b> ",
                    suffix_markup=question_suffix,
                )
                story.append(Spacer(1, 0.3*cm))

                PaperExporter._append_images_to_pdf_story(story, q.get("question_images"))

                if q_type in ['choice', 'multiple_choice'] and q.get('options'):
                    options = q.get('options', [])
                    for opt in options:
                        PaperExporter._add_text_with_latex_to_story(
                            story,
                            opt,
                            normal_style,
                            chinese_font,
                            prefix_markup="&nbsp;&nbsp;&nbsp;&nbsp;",
                        )

                if q.get('explanation'):
                    PaperExporter._add_text_with_latex_to_story(
                        story,
                        q["explanation"],
                        normal_style,
                        chinese_font,
                        prefix_markup="<b>解析：</b>",
                    )

                story.append(Spacer(1, 0.5*cm))
            
            # 如果包含答案，添加答案页
            if paper_data.get("answer_key"):
                story.append(PageBreak())
                story.append(Paragraph("<b>参考答案</b>", title_style))
                story.append(Spacer(1, 0.5*cm))
                
                answer_key = PaperExporter._normalize_answer_key(paper_data.get("answer_key"))
                for i, q in enumerate(questions, 1):
                    answer = answer_key.get(str(i), {}).get("answer", "")
                    PaperExporter._add_text_with_latex_to_story(
                        story,
                        str(answer) if answer is not None else "",
                        normal_style,
                        chinese_font,
                        prefix_markup=f"<b>{i}.</b> ",
                    )
                    story.append(Spacer(1, 0.3*cm))
                    PaperExporter._append_images_to_pdf_story(
                        story,
                        (q.get("answer_images") or []) + (q.get("solution_images") or []),
                        max_width=12 * cm,
                        max_height=7 * cm,
                    )
            
            # 生成PDF
            doc.build(story)
            logger.info(f"试卷PDF导出成功: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"导出PDF失败: {e}", exc_info=True)
            raise ValueError(f"导出PDF失败: {str(e)}")
    
    @staticmethod
    def export_to_word(paper_data: Dict[str, Any], output_path: str) -> str:
        """
        导出试卷为Word文档
        
        Args:
            paper_data: 试卷数据
            output_path: 输出文件路径
        
        Returns:
            str: 输出文件路径
        """
        if not DOCX_AVAILABLE:
            error_msg = DOCX_ERROR_MSG or "Word导出功能需要安装python-docx库。请运行: pip install python-docx"
            raise ValueError(error_msg)
        
        try:
            # 再次检查导入（运行时检查，处理安装了错误docx包的情况）
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError as e:
                raise ValueError(f"python-docx库导入失败: {e}。请运行: pip uninstall docx && pip install python-docx")
            
            # 创建Word文档
            doc = Document()
            
            # 设置中文字体
            # 获取默认样式并设置中文字体
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'  # 设置中文字体
            font.size = Pt(12)
            
            # 设置标题
            title = doc.add_heading(paper_data.get("title", "试卷"), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加试卷信息
            info_para = doc.add_paragraph()
            if paper_data.get("subject"):
                info_para.add_run(f"科目：{paper_data['subject']}  ")
            if paper_data.get("grade_level"):
                info_para.add_run(f"学段：{paper_data['grade_level']}  ")
            if paper_data.get("total_questions"):
                info_para.add_run(f"题数：{paper_data['total_questions']}  ")
            if paper_data.get("time_limit"):
                info_para.add_run(f"时长：{paper_data['time_limit']}分钟")
            
            doc.add_paragraph()  # 空行
            
            # 添加题目
            questions = PaperExporter._normalize_export_questions(paper_data.get("questions", []))
            for i, q in enumerate(questions, 1):
                # 获取题型标签
                q_type = q.get('type', '')
                type_label = ''
                if q_type == 'choice':
                    type_label = '单选题'
                elif q_type == 'multiple_choice':
                    type_label = '多选题'
                elif q_type == 'fill':
                    type_label = '填空题'
                elif q_type == 'judge':
                    type_label = '判断题'
                elif q_type == 'essay':
                    type_label = '简答题'
                elif q_type == 'calculation':
                    type_label = '计算题'

                # 题目
                question_para = doc.add_paragraph()
                question_para.add_run(f"{i}. ").bold = True

                # 处理题目文本，支持LaTeX公式
                question_text = q.get('question', '') or ''
                text_parts = PaperExporter._convert_latex_to_word_math(str(question_text))
                for part_text, is_math, image_path in text_parts:
                    if is_math and image_path:
                        # LaTeX公式：插入渲染的图片
                        try:
                            from docx.shared import Inches
                            question_para.add_run(" ")  # 添加空格
                            run = question_para.add_run()
                            run.add_picture(image_path, width=Inches(2))
                            question_para.add_run(" ")  # 添加空格
                            # 删除临时图片文件
                            try:
                                os.unlink(image_path)
                            except:
                                pass
                        except Exception as e:
                            logger.warning(f"插入公式图片失败，使用文本: {e}")
                            run = question_para.add_run(f" {part_text} ")
                            run.italic = True
                    elif is_math:
                        # LaTeX公式：使用斜体显示
                        run = question_para.add_run(f" {part_text} ")
                        run.italic = True
                    else:
                        question_para.add_run(part_text)

                # 添加题型和分数标注
                if type_label:
                    question_para.add_run(f"（{type_label}）").italic = True
                question_para.add_run(f"（{q.get('points', 5)}分）").italic = True

                # 显示选项（单选题和多选题）
                PaperExporter._append_images_to_word_doc(doc, q.get("question_images"))
                if q_type in ['choice', 'multiple_choice'] and q.get('options'):
                    for opt in q.get('options', []):
                        opt_para = doc.add_paragraph(style='List Bullet')
                        # 处理选项文本，支持LaTeX公式
                        opt_text = str(opt) if opt is not None else ''
                        opt_parts = PaperExporter._convert_latex_to_word_math(opt_text)
                        for part_text, is_math, image_path in opt_parts:
                            if is_math and image_path:
                                # LaTeX公式：插入渲染的图片
                                try:
                                    from docx.shared import Inches
                                    opt_para.add_run(" ")
                                    run = opt_para.add_run()
                                    run.add_picture(image_path, width=Inches(1.5))
                                    opt_para.add_run(" ")
                                    # 删除临时图片文件
                                    try:
                                        os.unlink(image_path)
                                    except:
                                        pass
                                except Exception as e:
                                    logger.warning(f"插入公式图片失败，使用文本: {e}")
                                    run = opt_para.add_run(f" {part_text} ")
                                    run.italic = True
                            elif is_math:
                                run = opt_para.add_run(f" {part_text} ")
                                run.italic = True
                            else:
                                opt_para.add_run(part_text)
                
                doc.add_paragraph()  # 空行
            
            # 如果包含答案，添加答案页
            if paper_data.get("answer_key"):
                doc.add_page_break()
                doc.add_heading("参考答案", 1)
                
                answer_key = PaperExporter._normalize_answer_key(paper_data.get("answer_key"))
                for i, q in enumerate(questions, 1):
                    answer = answer_key.get(str(i), {}).get("answer", "") if answer_key.get(str(i)) else ""
                    answer = str(answer) if answer is not None else ""
                    answer_para = doc.add_paragraph()
                    answer_para.add_run(f"{i}. ").bold = True
                    # 处理答案文本，支持LaTeX公式
                    answer_parts = PaperExporter._convert_latex_to_word_math(answer)
                    for part_text, is_math, image_path in answer_parts:
                        if is_math and image_path:
                            # LaTeX公式：插入渲染的图片
                            try:
                                from docx.shared import Inches
                                answer_para.add_run(" ")
                                run = answer_para.add_run()
                                run.add_picture(image_path, width=Inches(1.5))
                                answer_para.add_run(" ")
                                # 删除临时图片文件
                                try:
                                    os.unlink(image_path)
                                except:
                                    pass
                            except Exception as e:
                                logger.warning(f"插入公式图片失败，使用文本: {e}")
                                run = answer_para.add_run(f" {part_text} ")
                                run.italic = True
                        elif is_math:
                            run = answer_para.add_run(f" {part_text} ")
                            run.italic = True
                        else:
                            answer_para.add_run(part_text)
            
            # 保存文档
            if paper_data.get("answer_key"):
                for q in questions:
                    PaperExporter._append_images_to_word_doc(
                        doc,
                        (q.get("answer_images") or []) + (q.get("solution_images") or []),
                        width_inches=4.2,
                    )

            doc.save(output_path)
            logger.info(f"试卷Word导出成功: {output_path}")
            return output_path
            
        except ImportError as e:
            logger.error(f"python-docx导入失败: {e}", exc_info=True)
            raise ValueError("Word导出功能需要安装python-docx库：pip install python-docx")
        except Exception as e:
            logger.error(f"导出Word失败: {e}", exc_info=True)
            raise ValueError(f"导出Word失败: {str(e)}")

